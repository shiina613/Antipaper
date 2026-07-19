"""Deterministic document ingestion into the shared NormalizedDocument contract."""

from __future__ import annotations

from dataclasses import dataclass, field
import hashlib
from io import BytesIO
import os
from pathlib import Path
import re
import time
from typing import Iterable
import unicodedata

import fitz

from ..intelligence.contracts import Citation, DocumentChunk, NormalizedDocument


MAX_FILE_SIZE_BYTES = 25 * 1024 * 1024


def _default_table_budget_seconds() -> float:
    """Resolve the per-document table-detection time budget from the environment.

    `INGESTION_TABLE_BUDGET_SECONDS` sets the cap directly. The legacy on/off switch
    `INGESTION_EXTRACT_TABLES=0` maps to a zero budget (skip detection entirely).
    """
    if os.getenv("INGESTION_EXTRACT_TABLES", "1").strip().lower() in {"0", "false", "no"}:
        return 0.0
    return max(0.0, float(os.getenv("INGESTION_TABLE_BUDGET_SECONDS", "8.0")))


class IngestionError(RuntimeError):
    """Base error for deterministic document ingestion failures."""


class UnsupportedFileError(IngestionError):
    """Raised when the input type is not supported by the MVP."""


class FileTooLargeError(IngestionError):
    """Raised when the input exceeds the API contract size limit."""


@dataclass(frozen=True)
class IngestionOptions:
    """Runtime options for the document ingestor."""

    max_file_size_bytes: int = MAX_FILE_SIZE_BYTES
    max_pages: int | None = None
    table_extraction_budget_seconds: float = field(default_factory=_default_table_budget_seconds)


@dataclass(frozen=True)
class StitchedPage:
    """Native text and table content extracted from one document page."""

    page_number: int
    content: str


@dataclass
class _SectionState:
    chapter: str | None = None
    section: str | None = None
    article: str | None = None
    clause: str | None = None
    point: str | None = None


@dataclass
class _LogicalUnit:
    page: int
    text_parts: list[str]
    chapter: str | None = None
    section: str | None = None
    article: str | None = None
    clause: str | None = None
    point: str | None = None

    @property
    def text(self) -> str:
        return " ".join(part.strip() for part in self.text_parts if part.strip()).strip()


class DocumentIngestor:
    """Load PDF/DOCX files into `NormalizedDocument`.

    The ingestor is deterministic, does not call any LLM, and reads only native
    document content through PyMuPDF or python-docx.
    """

    _chapter_pattern = re.compile(r"^\s*(CHƯƠNG|CHUONG)\s+([IVXLCDM\d]+)\b.*", re.IGNORECASE)
    _section_pattern = re.compile(r"^\s*(MỤC|MUC)\s+(\d+[a-zA-Z]?)\b.*", re.IGNORECASE)
    _article_pattern = re.compile(r"^\s*(ĐIỀU|DIEU)\s+(\d+[a-zA-Z]?)\b[.:]?\s*(.*)", re.IGNORECASE)
    _clause_pattern = re.compile(r"^\s*((?:Khoản|Khoan)\s+\d+|\d+\.)\s+.*", re.IGNORECASE)
    _point_pattern = re.compile(r"^\s*([a-zA-ZđĐ])\)\s+.*")

    def __init__(self, options: IngestionOptions | None = None) -> None:
        self.options = options or IngestionOptions()

    def ingest(self, path: str | Path) -> NormalizedDocument:
        source_path = Path(path)
        self._validate_file(source_path)
        suffix = source_path.suffix.lower()

        if suffix == ".pdf":
            pages = self._load_pdf_pages(source_path)
        elif suffix == ".docx":
            pages = self._load_docx_pages(source_path)
        else:
            raise UnsupportedFileError("Chỉ hỗ trợ PDF hoặc DOCX.")

        chunks, citations = self._build_chunks_and_citations(pages)
        return NormalizedDocument(
            document_id=self._document_id(source_path),
            file_name=source_path.name,
            page_count=max((page.page_number for page in pages), default=1),
            chunks=chunks,
            citations=citations,
        )

    def ingest_bytes(
        self,
        *,
        file_name: str,
        file_bytes: bytes,
        document_id: str,
    ) -> tuple[NormalizedDocument, list[StitchedPage]]:
        """Ingest an upload directly from bytes without retaining a source file."""
        suffix = Path(file_name).suffix.lower()
        if len(file_bytes) > self.options.max_file_size_bytes:
            raise FileTooLargeError("File exceeds the 25 MB upload limit.")
        if suffix == ".pdf":
            pages = self._load_pdf_bytes(file_bytes)
        elif suffix == ".docx":
            pages = self._load_docx_bytes(file_bytes)
        else:
            raise UnsupportedFileError("Only PDF and DOCX files are supported.")
        chunks, citations = self._build_chunks_and_citations(pages)
        return (
            NormalizedDocument(
                document_id=document_id,
                file_name=file_name,
                page_count=max((page.page_number for page in pages), default=1),
                chunks=chunks,
                citations=citations,
            ),
            pages,
        )

    def _validate_file(self, path: Path) -> None:
        if not path.exists():
            raise FileNotFoundError(f"Document not found: {path}")
        if path.stat().st_size > self.options.max_file_size_bytes:
            raise FileTooLargeError("File vượt quá giới hạn 25 MB.")
        if path.suffix.lower() not in {".pdf", ".docx"}:
            raise UnsupportedFileError("Chỉ hỗ trợ PDF hoặc DOCX.")

    def _load_pdf_pages(self, path: Path) -> list[StitchedPage]:
        return self._load_pdf_bytes(path.read_bytes())

    def _load_pdf_bytes(self, file_bytes: bytes) -> list[StitchedPage]:
        pages: list[StitchedPage] = []
        # Native table detection runs a full per-page layout analysis and is the single most
        # expensive part of parsing, which dominates wall-clock on constrained (serverless) CPU
        # and, because it holds the GIL, multiplies under concurrent uploads. Set the budget to
        # 0 to skip it entirely; otherwise cap the *total* time spent on it per document so a
        # long document (or a slow CPU) can never let table detection blow the ingestion budget.
        # Text extraction — which produces every chunk and citation — is never affected.
        table_budget = self.options.table_extraction_budget_seconds
        table_time_spent = 0.0
        with fitz.open(stream=file_bytes, filetype="pdf") as document:
            page_limit = min(self.options.max_pages or document.page_count, document.page_count)
            for page_index in range(page_limit):
                page = document[page_index]
                text = page.get_text("text", sort=True).strip()
                table_markdown = ""
                if table_budget > 0 and table_time_spent < table_budget:
                    started = time.perf_counter()
                    table_markdown = self._extract_native_tables_markdown(page)
                    table_time_spent += time.perf_counter() - started
                content_parts = [part for part in [text, table_markdown] if part]
                pages.append(
                    StitchedPage(
                        page_number=page_index + 1,
                        content=self._normalize_unicode("\n\n".join(content_parts)),
                    )
                )
        return pages

    def _extract_native_tables_markdown(self, page: fitz.Page) -> str:
        try:
            tables = page.find_tables().tables
        except Exception:
            return ""

        markdown_tables: list[str] = []
        for table in tables:
            rows = table.extract()
            markdown = self._rows_to_markdown(rows)
            if markdown:
                markdown_tables.append(markdown)
        return "\n\n".join(markdown_tables)

    def _load_docx_pages(self, path: Path) -> list[StitchedPage]:
        return self._load_docx_bytes(path.read_bytes())

    def _load_docx_bytes(self, file_bytes: bytes) -> list[StitchedPage]:
        try:
            from docx import Document
        except ImportError as exc:
            raise IngestionError("python-docx is required for DOCX ingestion.") from exc

        document = Document(BytesIO(file_bytes))
        parts: list[str] = []
        parts.extend(paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip())
        for table in document.tables:
            rows = [[cell.text.strip() for cell in row.cells] for row in table.rows]
            markdown = self._rows_to_markdown(rows)
            if markdown:
                parts.append(markdown)
        return [StitchedPage(page_number=1, content="\n\n".join(parts))]

    def _build_chunks_and_citations(
        self,
        pages: Iterable[StitchedPage],
    ) -> tuple[list[DocumentChunk], dict[str, Citation]]:
        logical_units = self._build_logical_units(pages)
        chunks: list[DocumentChunk] = []
        citations: dict[str, Citation] = {}
        page_indices: dict[int, int] = {}

        for unit in logical_units:
            text = unit.text
            if not text:
                continue
            page_indices[unit.page] = page_indices.get(unit.page, 0) + 1
            chunk_id = f"P{unit.page}-D{page_indices[unit.page]}"
            excerpt = self._excerpt(text)
            chunk = DocumentChunk(
                chunk_id=chunk_id,
                page=unit.page,
                text=text,
                chapter=unit.chapter,
                section=unit.section,
                article=unit.article,
                clause=unit.clause,
                point=unit.point,
            )
            chunks.append(chunk)
            citations[chunk_id] = Citation(
                page=unit.page,
                chapter=unit.chapter,
                section=unit.section,
                article=unit.article,
                clause=unit.clause,
                point=unit.point,
                excerpt=excerpt,
            )

        return chunks, citations

    def _build_logical_units(self, pages: Iterable[StitchedPage]) -> list[_LogicalUnit]:
        units: list[_LogicalUnit] = []
        state = _SectionState()
        current: _LogicalUnit | None = None

        def flush() -> None:
            nonlocal current
            if current is not None and current.text:
                units.append(current)
            current = None

        for page in pages:
            for paragraph in self._split_content(page.content):
                boundary = self._update_section_state(paragraph, state)
                if current is not None and current.page != page.page_number:
                    flush()
                if boundary:
                    flush()
                if current is None:
                    current = _LogicalUnit(
                        page=page.page_number,
                        text_parts=[],
                        chapter=state.chapter,
                        section=state.section,
                        article=state.article,
                        clause=state.clause,
                        point=state.point,
                    )
                current.text_parts.append(paragraph)

        flush()
        return units

    def _update_section_state(self, text: str, state: _SectionState) -> bool:
        first_line = self._normalize_unicode(text).splitlines()[0].strip()
        chapter_match = self._chapter_pattern.match(first_line)
        if chapter_match:
            state.chapter = first_line
            state.section = None
            state.article = None
            state.clause = None
            state.point = None
            return True

        section_match = self._section_pattern.match(first_line)
        if section_match:
            state.section = first_line
            state.article = None
            state.clause = None
            state.point = None
            return True

        article_match = self._article_pattern.match(first_line)
        if article_match:
            state.article = f"Điều {article_match.group(2)}"
            state.clause = None
            state.point = None
            return True

        clause_match = self._clause_pattern.match(first_line)
        if clause_match:
            state.clause = clause_match.group(1).rstrip(".")
            state.point = None
            return True

        point_match = self._point_pattern.match(first_line)
        if point_match:
            state.point = point_match.group(1).lower()
            return True

        return False

    @staticmethod
    def _split_content(content: str) -> list[str]:
        normalized = DocumentIngestor._normalize_unicode(re.sub(r"\r\n?", "\n", content or ""))
        paragraphs = [
            re.sub(r"\s+", " ", paragraph).strip()
            for paragraph in re.split(r"\n{2,}", normalized)
            if paragraph.strip()
        ]
        if len(paragraphs) <= 1:
            paragraphs = [
                re.sub(r"\s+", " ", line).strip()
                for line in normalized.splitlines()
                if line.strip()
            ]
        return [paragraph for paragraph in paragraphs if len(paragraph) >= 3]

    @staticmethod
    def _normalize_unicode(value: str) -> str:
        return unicodedata.normalize("NFC", value or "")

    @staticmethod
    def _rows_to_markdown(rows: list[list[str | None]]) -> str:
        clean_rows = [
            [str(cell or "").strip().replace("\n", " ") for cell in row]
            for row in rows
            if any(str(cell or "").strip() for cell in row)
        ]
        if not clean_rows:
            return ""

        column_count = max(len(row) for row in clean_rows)
        padded_rows = [row + [""] * (column_count - len(row)) for row in clean_rows]
        header = padded_rows[0]
        separator = ["---"] * column_count
        body = padded_rows[1:]

        def format_row(row: list[str]) -> str:
            return "| " + " | ".join(row) + " |"

        return "\n".join(
            [format_row(header), format_row(separator), *[format_row(row) for row in body]]
        )

    @staticmethod
    def _excerpt(text: str, max_chars: int = 220) -> str:
        compact = re.sub(r"\s+", " ", text).strip()
        if len(compact) <= max_chars:
            return compact
        prefix = compact[: max_chars - 3].rstrip()
        if " " in prefix:
            prefix = prefix.rsplit(" ", 1)[0]
        return prefix + "..."

    @staticmethod
    def _document_id(path: Path) -> str:
        digest = hashlib.sha256(path.read_bytes()).hexdigest()
        return digest[:16]


def ingest_document(
    path: str | Path,
    options: IngestionOptions | None = None,
) -> NormalizedDocument:
    return DocumentIngestor(options).ingest(path)
