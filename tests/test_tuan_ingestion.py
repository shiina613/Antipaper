from __future__ import annotations

from pathlib import Path

import fitz
import pytest

from ingestion import (
    FileTooLargeError,
    IngestionError,
    IngestionOptions,
    UnsupportedFileError,
    ingest_document,
)
from intelligence import NormalizedDocument


def test_pdf_ingestion_returns_normalized_document_with_citations(tmp_path: Path) -> None:
    pdf_path = tmp_path / "demo.pdf"
    document = fitz.open()
    page_1 = document.new_page()
    page_1.insert_text(
        (72, 72),
        "CHUONG I\nDieu 1. Pham vi dieu chinh\n1. Co quan chu tri lap ke hoach trien khai.\nVan ban nay duoc dung de test chunking.",
    )
    page_2 = document.new_page()
    page_2.insert_text(
        (72, 72),
        "Dieu 2. Trach nhiem thuc hien\nKhoan 1 Cac don vi gui bao cao dinh ky.",
    )
    document.save(pdf_path)
    document.close()

    normalized = ingest_document(
        pdf_path,
        IngestionOptions(use_yolo_tables=False),
    )

    assert isinstance(normalized, NormalizedDocument)
    assert normalized.page_count == 2
    assert normalized.document_id
    assert normalized.file_name == "demo.pdf"
    assert normalized.chunks
    assert set(normalized.citations) == {chunk.chunk_id for chunk in normalized.chunks}
    assert {chunk.page for chunk in normalized.chunks} == {1, 2}
    assert any(chunk.chapter == "CHUONG I" for chunk in normalized.chunks)
    assert any(chunk.article == "Điều 1" for chunk in normalized.chunks)
    assert any(chunk.clause in {"1", "Khoản 1"} for chunk in normalized.chunks)
    assert any(
        "Co quan chu tri" in chunk.text and "Van ban nay" in chunk.text
        for chunk in normalized.chunks
    )


def test_docx_ingestion_smoke(tmp_path: Path) -> None:
    pytest.importorskip("docx")
    from docx import Document

    docx_path = tmp_path / "demo.docx"
    document = Document()
    document.add_paragraph("Chương I")
    document.add_paragraph("Điều 3. Nội dung chính")
    document.add_paragraph("1. Đơn vị phụ trách chịu trách nhiệm báo cáo.")
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Hạng mục"
    table.cell(0, 1).text = "Tiến độ"
    table.cell(1, 0).text = "Kinh phí"
    table.cell(1, 1).text = "Quý III"
    document.save(docx_path)

    normalized = ingest_document(docx_path)

    assert normalized.page_count == 1
    assert normalized.chunks
    assert any("| Hạng mục | Tiến độ |" in chunk.text for chunk in normalized.chunks)
    assert normalized.citation_whitelist == set(normalized.citations)


def test_ingestion_rejects_unsupported_and_large_files(tmp_path: Path) -> None:
    text_path = tmp_path / "demo.txt"
    text_path.write_text("unsupported", encoding="utf-8")
    with pytest.raises(UnsupportedFileError):
        ingest_document(text_path)

    pdf_path = tmp_path / "large.pdf"
    pdf_path.write_bytes(b"%PDF-1.7\n")
    with pytest.raises(FileTooLargeError):
        ingest_document(pdf_path, IngestionOptions(max_file_size_bytes=1))


def test_pdf_ingestion_can_require_yolo_weights(tmp_path: Path) -> None:
    pdf_path = tmp_path / "demo.pdf"
    document = fitz.open()
    document.new_page().insert_text((72, 72), "Điều 1. Nội dung")
    document.save(pdf_path)
    document.close()

    with pytest.raises(IngestionError, match="YOLOv8 weights"):
        ingest_document(
            pdf_path,
            IngestionOptions(
                use_yolo_tables=True,
                require_yolo_weights=True,
                yolo_model_path=tmp_path / "missing.pt",
            ),
        )
